# -*- coding: utf-8 -*-
"""Generate project report as Word document."""

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from pathlib import Path

OUTPUT = Path(__file__).resolve().parent.parent / "BAO_CAO_REACH_Church_App.docx"


def set_cell_shading(cell, color_hex: str):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    shading = OxmlElement("w:shd")
    shading.set(qn("w:fill"), color_hex)
    cell._tc.get_or_add_tcPr().append(shading)


def add_table(doc, headers, rows, header_color="48BCE1"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        set_cell_shading(hdr[i], header_color)
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(255, 255, 255)
                r.font.size = Pt(10)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            table.rows[ri + 1].cells[ci].text = str(val)
            for p in table.rows[ri + 1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    doc.add_paragraph()


def main():
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(13)

    # Title
    title = doc.add_heading("BÁO CÁO DỰ ÁN", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = doc.add_paragraph("R.E.A.C.H Church Vietnam — REACH Church App")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in sub.runs:
        r.bold = True
        r.font.size = Pt(14)
    doc.add_paragraph()

    # Meta table
    add_table(
        doc,
        ["Hạng mục", "Thông tin"],
        [
            ("Tên dự án", "R.E.A.C.H Church Vietnam (REACH Church App)"),
            ("Mã nguồn", "reach-church"),
            ("Repository", "github.com/truonga-dev/REACH_Church_App"),
            ("Người phát triển", "Truonga (atruong102005@gmail.com)"),
            ("Phiên bản", "0.1.0"),
            ("Ngày báo cáo", "31/05/2026"),
            ("Giai đoạn", "Phát triển chức năng — MVP đã có giao diện & backend cơ bản"),
        ],
    )

    sections = [
        (
            "1. Tóm tắt điều hành",
            [
                "R.E.A.C.H Church Vietnam là ứng dụng web mobile-first phục vụ Hội Thánh REACH, giúp tín hữu và khách thăm kết nối với đời sống hội thánh qua điện thoại hoặc trình duyệt. Ứng dụng tích hợp Supabase làm backend, hỗ trợ PWA (cài đặt như app native) và cung cấp trải nghiệm hoàn toàn bằng tiếng Việt.",
                "Dự án đã vượt qua giai đoạn khởi tạo, hiện có 9 màn hình người dùng, 1 trang quản trị, 1 API nội bộ, khoảng 44 file trong Git. Giao diện REACH Church đã được triển khai đầy đủ; dữ liệu động được đồng bộ qua Supabase với fallback dữ liệu tĩnh khi chưa có kết nối.",
            ],
        ),
        (
            "2. Bối cảnh & mục tiêu",
            [
                "2.1. Bối cảnh: Hội thánh cần kênh số để cập nhật bản tin, chia sẻ bài giảng, hỗ trợ đọc Kinh Thánh, tiếp nhận cầu nguyện và giới thiệu ban ngành. Trước đây thông tin chủ yếu qua Zalo, Facebook — thiếu nền tảng tập trung.",
                "2.2. Mục tiêu: (1) Nền tảng web mobile-first + PWA — Hoàn thành; (2) Trang chủ tổng hợp — Hoàn thành; (3) Kinh Thánh tiếng Việt — Hoàn thành; (4) Quản trị nội dung Admin — Hoàn thành; (5) Hồ sơ & cầu nguyện — Cơ bản; (6) Thư viện media đầy đủ — Đang phát triển; (7) Xác thực người dùng — Chưa triển khai.",
            ],
        ),
        (
            "3. Phạm vi dự án",
            [
                "Trong phạm vi: Ứng dụng web responsive; 9 trang người dùng + 1 admin; API Kinh Thánh; Supabase (news, sermons, prayers, profiles); PWA; Bottom navigation; YouTube embed bài giảng; Git/GitHub.",
                "Ngoài phạm vi: App iOS/Android native; Thanh toán trực tuyến; Push notification thật; Đa ngôn ngữ; Chat realtime.",
            ],
        ),
        (
            "4. Kiến trúc hệ thống",
            [
                "Mô hình Frontend-heavy: Next.js 16 App Router render phía client; Supabase làm BaaS; Kinh Thánh qua API Route đọc file JSON tĩnh (public/bible_vie.json).",
                "Luồng: Người dùng → Next.js (Trang chủ, Kinh Thánh, Thư viện, Admin...) → Supabase Cloud (news, sermons, prayers, profiles) và /api/bible (local JSON).",
            ],
        ),
    ]

    for heading, paragraphs in sections:
        doc.add_heading(heading, level=1)
        for p in paragraphs:
            doc.add_paragraph(p)

    doc.add_heading("5. Công nghệ sử dụng", level=1)
    add_table(
        doc,
        ["Hạng mục", "Công nghệ", "Phiên bản", "Vai trò"],
        [
            ("Framework", "Next.js (App Router)", "16.2.6", "SSR, routing, API"),
            ("UI Library", "React", "19.2.4", "Component UI"),
            ("Ngôn ngữ", "TypeScript", "5.x", "Type safety"),
            ("Backend/DB", "Supabase", "2.106.x", "Database, storage"),
            ("Icons", "Lucide React", "1.17.x", "Icon set"),
            ("Rich Text", "react-quill-new", "3.8.x", "Editor Admin"),
            ("PWA", "next-pwa", "5.6.x", "Cài như app"),
            ("Linting", "ESLint", "9.x", "Code quality"),
            ("VCS", "Git + GitHub", "—", "Quản lý mã nguồn"),
        ],
    )

    doc.add_heading("6. Chức năng chi tiết", level=1)
    add_table(
        doc,
        ["Trang", "Đường dẫn", "Mô tả"],
        [
            ("Trang chủ", "/", "Bản tin, sự kiện, bài giảng YouTube, dưỡng linh, thông báo"),
            ("Kinh Thánh", "/bible", "66 sách KT tiếng Việt, API nội bộ"),
            ("Thư viện", "/library", "PDF từ Supabase; sách nói/video đang phát triển"),
            ("Mục vụ", "/ministry", "6 ban ngành, modal chi tiết"),
            ("Hồ sơ", "/profile", "Thông tin, cầu nguyện, quyên góp"),
            ("Cầu nguyện", "/prayer", "Form gửi nhu cầu cầu nguyện"),
            ("Dưỡng linh", "/devotional", "Bài đọc chi tiết, chia sẻ"),
            ("Tin tức", "/news/[id]", "Chi tiết bản tin từ Supabase"),
            ("Admin", "/admin", "CRUD bài giảng, tin, cầu nguyện, profiles"),
        ],
    )

    doc.add_heading("7. Cơ sở dữ liệu Supabase", level=1)
    add_table(
        doc,
        ["Bảng", "Mục đích", "Thao tác"],
        [
            ("news", "Bản tin, sự kiện, PDF, audio", "SELECT, INSERT, UPDATE, DELETE"),
            ("sermons", "Bài giảng + YouTube URL", "SELECT, INSERT, UPDATE, DELETE"),
            ("prayers", "Đề mục cầu nguyện", "SELECT, UPDATE, DELETE"),
            ("profiles", "Hồ sơ tín hữu", "SELECT, UPDATE"),
        ],
    )
    doc.add_paragraph(
        "Biến môi trường bắt buộc: NEXT_PUBLIC_SUPABASE_URL, NEXT_PUBLIC_SUPABASE_ANON_KEY"
    )

    doc.add_heading("8. Thiết kế giao diện", level=1)
    add_table(
        doc,
        ["Token", "Giá trị", "Ý nghĩa"],
        [
            ("Primary", "#48BCE1", "REACH Blue"),
            ("Secondary", "#F4CC30", "REACH Yellow"),
            ("Accent", "#F12D5C", "REACH Red/Pink"),
        ],
    )
    doc.add_paragraph(
        "UX: Mobile-first, Bottom Navigation 5 tab, dark mode tự động, PWA manifest, toast notifications, modal video YouTube. Ngôn ngữ: tiếng Việt (lang=vi)."
    )

    doc.add_heading("9. Tiến độ thực hiện", level=1)
    add_table(
        doc,
        ["Ngày", "Commit", "Nội dung"],
        [
            ("30/05/2026", "dcf6552", "Khởi tạo Next.js template"),
            ("31/05/2026", "361a1dd*", "Phát triển toàn bộ REACH Church App"),
            ("31/05/2026", "90f7ab6", "Khôi phục code sau git reset"),
            ("31/05/2026", "20b9f02", "Cập nhật README tiếng Việt"),
        ],
    )
    doc.add_paragraph("* Commit gốc, đã khôi phục qua 90f7ab6")

    add_table(
        doc,
        ["Module", "UI", "Logic", "Backend", "Tổng"],
        [
            ("Trang chủ", "✅", "✅", "✅", "90%"),
            ("Kinh Thánh", "✅", "✅", "✅", "100%"),
            ("Thư viện", "✅", "⚠️", "⚠️", "40%"),
            ("Mục vụ", "✅", "✅", "—", "80%"),
            ("Hồ sơ", "✅", "✅", "⚠️", "70%"),
            ("Cầu nguyện", "✅", "⚠️", "⚠️", "50%"),
            ("Dưỡng linh", "✅", "✅", "—", "80%"),
            ("Admin", "✅", "✅", "✅", "85%"),
            ("PWA", "✅", "✅", "—", "90%"),
            ("Auth", "❌", "❌", "❌", "0%"),
        ],
    )

    doc.add_heading("10. Quản lý mã nguồn & bảo mật", level=1)
    doc.add_paragraph(
        "GitHub: github.com/truonga-dev/REACH_Church_App, branch main, tài khoản truonga-dev. "
        ".env.local không push; .gitignore đầy đủ. "
        "Sự cố đã xử lý: git reset mất code (khôi phục từ reflog); sai credential GitHub (chuyển sang truonga-dev)."
    )
    add_table(
        doc,
        ["Hạng mục", "Trạng thái", "Ghi chú"],
        [
            (".env.local", "✅", "Keys không lên GitHub"),
            (".gitignore", "✅", "Bảo vệ secrets"),
            ("Admin auth", "⚠️", "Hardcode password — cần cải thiện"),
            ("Supabase RLS", "⚠️", "Cần kiểm tra policies"),
        ],
    )

    doc.add_heading("11. Quy trình phát triển & triển khai", level=1)
    doc.add_paragraph(
        "Cài đặt: git clone → npm install --legacy-peer-deps → cp .env.example .env.local → npm run dev → http://localhost:3000"
    )
    add_table(
        doc,
        ["Lệnh", "Mô tả"],
        [
            ("npm run dev", "Dev server 0.0.0.0:3000"),
            ("npm run build", "Build production"),
            ("npm run start", "Chạy production local"),
            ("npm run lint", "ESLint"),
        ],
    )
    doc.add_paragraph("Triển khai: Vercel — push GitHub, import repo, thêm env Supabase, deploy.")

    doc.add_heading("12. Rủi ro & hạn chế", level=1)
    add_table(
        doc,
        ["Rủi ro", "Mức", "Giải pháp"],
        [
            ("Admin password hardcode", "Cao", "Supabase Auth + RBAC"),
            ("Không user login", "TB", "Supabase Auth"),
            ("Thư viện chưa xong", "TB", "Storage + media API"),
            ("Form cầu nguyện chưa lưu DB", "TB", "INSERT prayers"),
            ("react-quill vs React 19", "Thấp", "--legacy-peer-deps"),
        ],
    )

    doc.add_heading("13. Kế hoạch phát triển tiếp theo", level=1)
    doc.add_paragraph("Giai đoạn A (1–2 tuần): Form cầu nguyện → DB; hoàn thiện thư viện; Supabase Auth; bảo mật Admin.")
    doc.add_paragraph("Giai đoạn B (2–4 tuần): Push notification; lịch đọc KT; quyên góp; SEO.")
    doc.add_paragraph("Giai đoạn C (4–6 tuần): Deploy Vercel + domain; RLS; testing; monitoring.")

    doc.add_heading("14. Kết luận & đánh giá", level=1)
    doc.add_paragraph(
        "R.E.A.C.H Church Vietnam có nền tảng kỹ thuật vững, giao diện hoàn chỉnh, tích hợp Supabase. "
        "Sẵn sàng hoàn thiện MVP và triển khai thử nghiệm. Ưu tiên: bảo mật Admin, Auth, form cầu nguyện, thư viện media."
    )
    add_table(
        doc,
        ["Tiêu chí", "Điểm", "Nhận xét"],
        [
            ("Kiến trúc kỹ thuật", "5/5", "Next.js 16 + Supabase + PWA"),
            ("Giao diện", "5/5", "Mobile-first, REACH branding"),
            ("Tính năng", "4/5", "9 trang + admin, một số chưa xong"),
            ("Backend", "3/5", "Thiếu Auth & RLS"),
            ("Bảo mật", "2/5", "Admin hardcode — cần sửa"),
            ("Tài liệu", "4/5", "README tiếng Việt"),
            ("Git", "4/5", "GitHub ổn định"),
        ],
    )

    doc.add_paragraph()
    footer = doc.add_paragraph("— Hết báo cáo —")
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note = doc.add_paragraph("Báo cáo được lập dựa trên trạng thái dự án tại 31/05/2026.")
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in note.runs:
        r.italic = True
        r.font.size = Pt(11)

    doc.save(OUTPUT)
    print(f"Created: {OUTPUT}")


if __name__ == "__main__":
    main()
